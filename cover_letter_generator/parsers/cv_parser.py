"""
CV Parser - Extracts and structures information from .docx CV files.
"""

import re
from pathlib import Path
from typing import List, Dict, Tuple
from docx import Document

from ..models import CVData, PersonalInfo, Experience, Education, Project


class CVParser:
    """Parser for extracting structured data from CV documents."""

    # Section headers to look for (case-insensitive)
    SECTION_KEYWORDS = {
        'experience': ['experience', 'work history', 'employment', 'professional experience', 'work experience'],
        'education': ['education', 'academic background', 'qualifications'],
        'skills': ['skills', 'technical skills', 'core competencies', 'expertise'],
        'projects': ['projects', 'personal projects', 'portfolio'],
        'summary': ['summary', 'profile', 'objective', 'about me'],
        'certifications': ['certifications', 'certificates', 'licenses']
    }

    def __init__(self, docx_path: str):
        """Initialize parser with path to .docx file."""
        self.docx_path = Path(docx_path)
        self.document = None
        self.paragraphs = []
        self.full_text = ""

    def extract_text_from_docx(self) -> str:
        """Extract all text from .docx file."""
        self.document = Document(self.docx_path)
        self.paragraphs = [p.text.strip() for p in self.document.paragraphs if p.text.strip()]
        self.full_text = '\n'.join(self.paragraphs)
        return self.full_text

    def extract_personal_info(self) -> PersonalInfo:
        """Extract personal information from the first few paragraphs."""
        personal_info = PersonalInfo()

        # Check first 5 paragraphs for personal info
        first_lines = self.full_text.split('\n')[:5]
        text_block = ' '.join(first_lines)

        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text_block)
        if emails:
            personal_info.email = emails[0]

        # Extract phone
        phone_pattern = r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,5}[-\s\.]?[0-9]{1,5}'
        phones = re.findall(phone_pattern, text_block)
        if phones:
            personal_info.phone = phones[0]

        # Extract LinkedIn
        if 'linkedin.com' in text_block.lower():
            linkedin_pattern = r'(https?://)?([a-z]{2,3}\.)?linkedin\.com/[^\s]+'
            linkedins = re.findall(linkedin_pattern, text_block, re.IGNORECASE)
            if linkedins:
                personal_info.linkedin = linkedins[0][1] if isinstance(linkedins[0], tuple) else linkedins[0]

        # Extract GitHub
        if 'github.com' in text_block.lower():
            github_pattern = r'(https?://)?github\.com/[^\s]+'
            githubs = re.findall(github_pattern, text_block, re.IGNORECASE)
            if githubs:
                personal_info.github = githubs[0]

        # Assume first line is the name if it doesn't contain email or phone
        if self.paragraphs and '@' not in self.paragraphs[0] and not re.search(phone_pattern, self.paragraphs[0]):
            personal_info.name = self.paragraphs[0]

        return personal_info

    def identify_sections(self) -> Dict[str, List[str]]:
        """Identify sections in the CV and map them to content."""
        sections = {}
        current_section = 'header'
        current_content = []

        for para in self.paragraphs:
            # Check if this paragraph is a section header
            section_found = False
            para_lower = para.lower().strip()

            for section_name, keywords in self.SECTION_KEYWORDS.items():
                for keyword in keywords:
                    if para_lower == keyword or (len(para.split()) <= 3 and keyword in para_lower):
                        # Save previous section
                        if current_content:
                            sections[current_section] = current_content
                        # Start new section
                        current_section = section_name
                        current_content = []
                        section_found = True
                        break
                if section_found:
                    break

            if not section_found:
                current_content.append(para)

        # Save last section
        if current_content:
            sections[current_section] = current_content

        return sections

    def parse_experience_section(self, content: List[str]) -> List[Experience]:
        """Parse experience entries from content."""
        experiences = []
        current_exp = None

        for line in content:
            # Try to detect company/role lines (usually short, may contain dates)
            if len(line.split()) <= 15 and not line.startswith(('•', '-', '–', '·')):
                # This might be a new experience entry
                if current_exp and (current_exp.company or current_exp.role):
                    experiences.append(current_exp)

                # Parse dates if present
                date_pattern = r'(\d{4}|present|current)'
                dates = re.findall(date_pattern, line, re.IGNORECASE)

                current_exp = Experience(
                    company=line.split('|')[0].split('–')[0].split('-')[0].strip() if '|' in line or '–' in line or '-' in line else line,
                    role=line.split('|')[1].strip() if '|' in line else '',
                    description='',
                    start_date=dates[0] if len(dates) > 0 else None,
                    end_date=dates[1] if len(dates) > 1 else None
                )
            elif current_exp:
                # This is a description or bullet point
                if line.startswith(('•', '-', '–', '·', '*')):
                    # Bullet point
                    bullet_text = line.lstrip('•-–·* ').strip()
                    current_exp.bullets.append(bullet_text)
                else:
                    # Regular description
                    if current_exp.description:
                        current_exp.description += ' ' + line
                    else:
                        current_exp.description = line

        # Add last experience
        if current_exp and (current_exp.company or current_exp.role):
            experiences.append(current_exp)

        return experiences

    def parse_education_section(self, content: List[str]) -> List[Education]:
        """Parse education entries from content."""
        education = []

        for i, line in enumerate(content):
            if len(line.split()) <= 15 and not line.startswith(('•', '-', '–', '·')):
                # Likely an education entry
                date_pattern = r'(\d{4})'
                dates = re.findall(date_pattern, line)

                # Try to extract degree, institution
                parts = line.split('|') if '|' in line else [line]

                edu = Education(
                    institution=parts[0].strip(),
                    degree=parts[1].strip() if len(parts) > 1 else line,
                    start_date=dates[0] if len(dates) > 0 else None,
                    end_date=dates[1] if len(dates) > 1 else None
                )

                # Check next line for additional details
                if i + 1 < len(content) and content[i+1].startswith(('•', '-', 'GPA')):
                    edu.description = content[i+1].lstrip('•-–·* ').strip()

                education.append(edu)

        return education

    def parse_skills_section(self, content: List[str]) -> List[str]:
        """Parse skills from content."""
        skills = []

        for line in content:
            # Remove bullet points
            line = line.lstrip('•-–·* ').strip()

            # Skills might be comma-separated or one per line
            if ',' in line:
                skills.extend([s.strip() for s in line.split(',') if s.strip()])
            elif '|' in line:
                skills.extend([s.strip() for s in line.split('|') if s.strip()])
            elif line:
                skills.append(line)

        return skills

    def parse_projects_section(self, content: List[str]) -> List[Project]:
        """Parse project entries from content."""
        projects = []
        current_project = None

        for line in content:
            # Short lines without bullets are likely project titles
            if len(line.split()) <= 10 and not line.startswith(('•', '-', '–', '·')):
                if current_project:
                    projects.append(current_project)

                current_project = Project(
                    name=line.split('|')[0].split('–')[0].strip(),
                    description=''
                )
            elif current_project:
                if line.startswith(('•', '-', '–', '·')):
                    current_project.highlights.append(line.lstrip('•-–·* ').strip())
                else:
                    if current_project.description:
                        current_project.description += ' ' + line
                    else:
                        current_project.description = line

        if current_project:
            projects.append(current_project)

        return projects

    def parse_cv(self) -> CVData:
        """Main parsing function - returns structured CV data."""
        # Extract text
        self.extract_text_from_docx()

        # Extract personal info
        personal_info = self.extract_personal_info()

        # Identify sections
        sections = self.identify_sections()

        # Parse each section
        experience = []
        if 'experience' in sections:
            experience = self.parse_experience_section(sections['experience'])

        education = []
        if 'education' in sections:
            education = self.parse_education_section(sections['education'])

        skills = []
        if 'skills' in sections:
            skills = self.parse_skills_section(sections['skills'])

        projects = []
        if 'projects' in sections:
            projects = self.parse_projects_section(sections['projects'])

        summary = None
        if 'summary' in sections:
            summary = ' '.join(sections['summary'])

        certifications = []
        if 'certifications' in sections:
            certifications = sections['certifications']

        # Create CVData object
        cv_data = CVData(
            personal_info=personal_info,
            summary=summary,
            experience=experience,
            education=education,
            skills=skills,
            projects=projects,
            certifications=certifications,
            raw_text=self.full_text
        )

        return cv_data


def parse_cv(docx_path: str) -> CVData:
    """
    Parse a .docx CV file and return structured data.

    Args:
        docx_path: Path to the .docx CV file

    Returns:
        CVData object with structured CV information
    """
    parser = CVParser(docx_path)
    return parser.parse_cv()
